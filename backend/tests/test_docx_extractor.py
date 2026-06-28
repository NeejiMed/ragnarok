from pathlib import Path

from backend.app.ingestion.extractors.docx_extractor import extract_docx


def make_text_docx(path: str, paragraphs: list[str]) -> None:
    """Generates a synthetic DOCX with one paragraph per string in `paragraphs`."""
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def test_plain_paragraph_docx_extraction():
    docx_path = Path(__file__).parent / "fixtures" / "basic-text.docx"
    make_text_docx(str(docx_path), ["Hello, World!", "This is a test."])
    extracted_content = extract_docx(str(docx_path))

    assert len(extracted_content) == 1
    assert extracted_content[0]["extraction_method"] == "native"
    assert extracted_content[0]["page"] is None


def test_docx_with_tables_extraction():
    docx_path = Path(__file__).parent / "fixtures" / "table-text.docx"
    from docx import Document

    doc = Document()
    doc.add_paragraph("This is a paragraph before the table.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Row 1 Col 1"
    table.cell(1, 1).text = "Row 1 Col 2"
    doc.save(str(docx_path))

    extracted_content = extract_docx(str(docx_path))

    assert len(extracted_content) == 1
    assert extracted_content[0]["extraction_method"] == "native"
    assert extracted_content[0]["page"] is None
    assert "This is a paragraph before the table." in extracted_content[0]["content"]
    assert "Header 1 | Header 2" in extracted_content[0]["content"]
    assert "Row 1 Col 1 | Row 1 Col 2" in extracted_content[0]["content"]
