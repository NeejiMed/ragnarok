from docx import Document as DocxDocument


def extract_docx(file_path: str) -> list[dict]:
    """
    Extracts text from a DOCX file: paragraphs + tables.
    DOCX has no reliable page concept, so we return to a single combined unit.
    """
    docx = DocxDocument(file_path)

    paragraph_text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())

    table_text = _extract_tables(docx)

    full_text = paragraph_text
    if table_text:
        full_text += "\n\n" + table_text

    return [{"content": full_text.strip(), "page": None, "extraction_method": "native"}]


def _extract_tables(docx: "DocxDocument") -> str:  # type: ignore
    """Extracts all tables as pipe-delimited rows, tables separated by blank lines."""
    table_blocks = []
    for table in docx.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_blocks.append("\n".join(rows))
    return "\n\n".join(table_blocks)
